from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_doc(sum=20000, period=5):


    # Create a new document
    doc = Document()

    # Add a header with the organization name
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Организация Банк"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the content of the document
    doc.add_paragraph("Кредитный документ:")
    doc.add_paragraph("Сумма кредита: " + sum + " рублей")
    doc.add_paragraph("Процентная ставка: 5% годовых")
    doc.add_paragraph("Срок кредита: " + period + " лет")

    # Calculation of monthly payments
    sum = (float(sum))
    period = int(period)
    loan_amount = sum
    interest_rate = 0.05
    loan_term_years = period
    monthly_interest_rate = interest_rate / 12
    loan_term_months = loan_term_years * 12

    monthly_payment = (loan_amount * monthly_interest_rate) / (1 - (1 + monthly_interest_rate) ** (-loan_term_months))

    # Add a table for the monthly payments
    table = doc.add_table(rows=loan_term_months + 1, cols=3)
    table.cell(0, 0).text = "Месяц"
    table.cell(0, 1).text = "Сумма платежа"
    table.cell(0, 2).text = "Остаток долга"

    for month in range(loan_term_months):
        payment_number = month + 1
        remaining_loan_amount = loan_amount * (1 + monthly_interest_rate) ** payment_number - monthly_payment * (
                    (1 + monthly_interest_rate) ** payment_number - 1) / monthly_interest_rate

        table.cell(month + 1, 0).text = str(payment_number)
        table.cell(month + 1, 1).text = str(round(monthly_payment, 2))
        table.cell(month + 1, 2).text = str(round(remaining_loan_amount, 2))
    path = "Кредитный документ.docx"
    doc.save(path)
    # Сохраняем документ
    return path
